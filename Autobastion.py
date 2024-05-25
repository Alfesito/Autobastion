from docx import Document
from openpyxl import Workbook
from tqdm import tqdm
import PyPDF2
import re
# Instalación de librerias: $ pip install python-docx openpyxl tqdm PyPDF2
 
def extract_titles(doc_path, headings):
    doc = Document(doc_path)
    titles = []
    domain = None
    current_subdomain1 = None
    current_subdomain2 = None
    current_control = None
    remediation = None
    verification = None
    impact = None
    is_sub2 = False
 
    for para in doc.paragraphs:
 
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split(' ')[1])
            text = para.text.strip()
 
            if level == 2:  # Dominio
                is_sub2 = False
                domain = text
                current_subdomain1 = None
                current_subdomain2 = None
                current_control = None
                remediation = None
                verification = None
                impact = None
            elif level == 4 and not is_sub2:  # Subdominio 1
                current_subdomain1 = text
                current_subdomain2 = None
                current_control = None
                remediation = None
                verification = None
                impact = None
                is_sub2 = True
            elif level == 4 and is_sub2:  # Subdominio 2
                current_subdomain2 = text
                current_control = None
                remediation = None
                verification = None
                impact = None
            elif level == 3:  # Control
                current_control = text
                # Extraer el número de control
                for key, value in headings.items():
                    value_words = value.split()[:6]
                    control_words = str(current_control).split()[:6]
                    if len(value_words) >= 6 and len(control_words) >= 6:
                        if value_words == control_words:  
                            control_number = key
                            break  # Salir del bucle una vez que se encuentra una coincidencia
                        else:
                            control_number = None
                titles.append((domain, current_subdomain1, current_subdomain2, control_number, current_control, remediation, verification, impact))
 
    return titles
 
def extract_text_sections(doc_path, section_title):
    doc = Document(doc_path)
    texts = []
    section_text = ""
    in_section = False
    exclude_phrases = ["CIS Controls:", "MITRE ATT&CK Mappings:"]  # Lista de frases a excluir
 
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and in_section and para.text not in exclude_phrases:
            texts.append(section_text.strip())
            in_section = False
            section_text = ""
        if section_title in para.text:
            in_section = True
        elif in_section and not any(exclude in para.text for exclude in exclude_phrases):
            section_text += para.text.strip() + "\n"
 
    if in_section:
        texts.append(section_text.strip())
 
    return texts
 
def extract_numbered_headings(pdf_path):
    headings = {}
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text = page.extract_text()  # Extraemos el texto de la página
            lines = text.split('\n')  # Dividimos el texto en líneas
            for line in lines:
                # Utilizamos una expresión regular para buscar cualquier cadena que comience con un número seguido de un punto
                matches = re.findall(r'(\d+\..*)', line)
                if matches:
                    for match in matches:
                        # Dividimos la cadena coincidente en ID y texto
                        parts = match.split(' ', 1)
                        # Añadimos el ID y el texto limpio al diccionario
                        headings[parts[0]] = parts[1].strip()
    return headings
 
def write_titles_to_excel(titles, remediation_texts, verification_texts, impact_texts, excel_path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Controls"
 
    headers = ["Dominio", "Subdominio1", "Subdominio2", "ID", "Control", "Remediación", "Verificación", "Impacto"]
    ws.append(headers)
 
    # Crear los títulos extendidos con textos de remediación, verificación e impacto
    extended_titles = []
    for i in range(len(titles)):
        extended_title = list(titles[i])  # Convertir tupla a lista para poder modificarla
        if i < len(remediation_texts):
            extended_title[5] = remediation_texts[i]
        if i < len(verification_texts):
            extended_title[6] = verification_texts[i]
        if i < len(impact_texts):
            extended_title[7] = impact_texts[i]
        extended_titles.append(extended_title)
 
    # Añadir cada título extendido a la hoja de cálculo
    for i, extended_title in enumerate(tqdm(extended_titles, desc="Copiando datos a Excel", unit="fila")):
        ws.append(extended_title)
 
    wb.save(excel_path)
 
def main():
    word_path = r'/Users/andresalfarofernandez/DocumentosPC/VisualStudio_code/Hacking/Autobastion/Templates/CIS_Debian_Linux_10_Benchmark_v2.0.0.docx'
    pdf_path = r'/Users/andresalfarofernandez/DocumentosPC/VisualStudio_code/Hacking/Autobastion/Templates/CIS_Debian_Linux_10_Benchmark_v2.0.0.pdf'
    excel_path = r'/Users/andresalfarofernandez/DocumentosPC/VisualStudio_code/Hacking/Autobastion/output.xlsx'
 
    # Barra de progreso principal para todo el proceso
    with tqdm(total=100, desc="Procesando documento de Word a Excel", unit="porcentaje") as pbar:
        headings = extract_numbered_headings(pdf_path)
        pbar.update(20)
        titles = extract_titles(word_path, headings)
        pbar.update(10)
 
        pbar.set_description("Extrayendo textos de remediación")
        remediation_texts = extract_text_sections(word_path, "Remediation:")
        pbar.update(20)
 
        pbar.set_description("Extrayendo textos de verificación")
        verification_texts = extract_text_sections(word_path, "Audit:")
        pbar.update(20)
 
        pbar.set_description("Extrayendo textos de impacto")
        impact_texts = extract_text_sections(word_path, "Description:")  # Cambiado a "Description:"
        pbar.update(20)
 
        pbar.set_description("Escribiendo datos en el archivo de Excel")
        write_titles_to_excel(titles, remediation_texts, verification_texts, impact_texts, excel_path)
        pbar.update(10)
 
    print(f"Datos copiados y pegados en {excel_path}")
 
if __name__ == "__main__":
    main()