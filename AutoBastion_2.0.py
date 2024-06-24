from docx import Document
from openpyxl import Workbook
from tqdm import tqdm
import PyPDF2

def extract_controls(doc_path, headings):
    doc = Document(doc_path)
    titles = []

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split(' ')[1])
            text = para.text.strip()
            if level == 3:  # Control, cambien el numero según el heading <-------------
                current_control = text.split(' (')[0]  # Elimina todo lo que hay a la derecha de un ' ('
                # Extraer el número de control
                for key, value in headings.items():
                    value_words = str(value)
                    control_words = str(current_control)
                    if value_words == control_words:
                        control_number = key
                        break  # Salir del bucle una vez que se encuentra una coincidencia
                    # Ajustar la longitud de control_words si es mayor que value_words
                    elif len(control_words) > len(value_words):
                        control_words = control_words.ljust(len(value_words))
                        if control_words == value_words:
                            control_number = key
                            break
                    else:
                        control_number = None
                titles.append((None, None, None, control_number, current_control, None, None, None, None))

    return titles

def extract_domains(headings, titles):
    domains = []
    for title in titles:
        domain = None
        subdomain1 = None
        subdomain2 = None
        control_number = title[3]
        current_control = title[4]

        if control_number:
            parts = control_number.split('.')
            if len(parts) > 0:
                domain = headings.get(parts[0], None)
                if domain == current_control:
                    domain = None
            if len(parts) > 1:
                subdomain1 = headings.get('.'.join(parts[:2]), None)
                if subdomain1 == current_control:
                    subdomain1 = None
            if len(parts) > 2:
                subdomain2 = headings.get('.'.join(parts[:3]), None)
                if subdomain2 == current_control:
                    subdomain2 = None

        domains.append((domain, subdomain1, subdomain2, control_number, current_control, None, None, None, None))

    return domains

def extract_text_sections_pdf(doc_path, section_title):
    texts = []
    section_text = ""
    in_section = False
    exclude_phrases = ["CIS Controls:", "MITRE ATT&CK Mappings:", "References:", "Additional Information:", "Default Value:"]  # Lista de frases a excluir

    reader = PyPDF2.PdfReader(doc_path)
    num_pages = len(reader.pages)

    for page_num in range(num_pages):
        page = reader.pages[page_num]
        text = page.extract_text()

        for line in text.splitlines():
            if in_section:
                if any(line.startswith(heading) for heading in exclude_phrases):
                    texts.append(section_text.strip())
                    in_section = False
                    section_text = ""
                else:
                    section_text += line + "\n"
            if section_title in line:
                in_section = True

        # Al final de la página, si estamos en la sección, continuamos
        if in_section and page_num == num_pages - 1:
            texts.append(section_text.strip())

    # Añadir el último texto de sección si quedó algo sin añadir
    if section_text.strip():
        texts.append(section_text.strip())

    return texts

def extract_text_sections_word(doc_path, section_title):
    doc = Document(doc_path)
    texts = []
    section_text = ""
    in_section = False
    in_default_value = False
    exclude_phrases = ["CIS Controls:", "MITRE ATT&CK Mappings:", "Audit:", "Remediation:", "References:", "Additional Information:", "Default Value:"]  # Lista de frases a excluir
    
    def process_paragraph(para):
        nonlocal section_text
        if para.style.name == 'List Paragraph':
            lines = para.text.split(' If ')
            for i, line in enumerate(lines):
                if i == 0:
                    section_text += '- ' + line.strip() + "\n"
                else:
                    section_text += "\nIf " + line.strip() + "\n"
        else:
            section_text += para.text.strip() + "\n"

    for para in doc.paragraphs:
        if section_title != "Default Value:":
            if para.style.name.startswith('Heading') and para.text in exclude_phrases and in_section: 
                texts.append(section_text.strip())
                in_section = False
                section_text = ""
            if para.style.name.startswith('Heading') and section_title == para.text:
                in_section = True
            elif in_section and not any(exclude in para.text for exclude in exclude_phrases):
                process_paragraph(para)
            elif para.text in exclude_phrases and in_section:
                texts.append(section_text.strip())
                in_section = False
                section_text = ""
        else:
            if in_default_value and para.text in exclude_phrases and para.text != "Default Value:":
                if not in_section:
                    section_text = "-"
                texts.append(section_text.strip())
                in_section = False
                in_default_value = False
                section_text = ""
            if para.style.name.startswith('Heading') and para.text == "Remediation:":
                in_default_value = True
            elif para.style.name.startswith('Heading') and section_title == para.text and in_default_value:
                in_section = True
            elif in_section:
                process_paragraph(para)
    if in_section:
        texts.append(section_text.strip())
 
    return texts


def extract_numbered_headings(word_path):
    headings = {}
    doc = Document(word_path)
    
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:  # Asegúrate de que haya al menos dos columnas
                key = cells[0].text.strip()
                value = cells[1].text.strip()
                clean_value = value.split(' (')[0].split(' .')[0].strip()
                headings[key] = clean_value
        
    return headings

def write_titles_to_excel(titles, remediation_texts, default_value_texts, verification_texts, impact_texts, excel_path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Controls"
    headers = ["Dominio", "Subdominio1", "Subdominio2", "ID", "Control", "Remediación", "Valor por defecto", "Verificación", "Impacto"]
    ws.append(headers)
    extended_titles = []

    for i in range(len(titles)):
        extended_title = list(titles[i])  # Convertir la tupla a lista
        if i < len(remediation_texts):
            extended_title[5] = remediation_texts[i]
        if i < len(default_value_texts):
            extended_title[6] = default_value_texts[i]
        if i < len(verification_texts):
            extended_title[7] = verification_texts[i]
        if i < len(impact_texts):
            extended_title[8] = impact_texts[i]
        extended_titles.append(extended_title)

    for i, extended_title in enumerate(tqdm(extended_titles, desc="Copiando datos a Excel", unit="fila")):
        ws.append(extended_title)
    
    merge_consecutive_rows(ws)
    wb.save(excel_path)


def merge_consecutive_rows(ws):
    max_row = ws.max_row
    for col in [1,2,3]:
        start_row = 2
        while start_row <= max_row:
            end_row = start_row
            while end_row < max_row and ws.cell(row=end_row, column=col).value == ws.cell(row=end_row + 1, column=col).value and ws.cell(row=end_row + 1, column=col).value != None:
                end_row += 1
            if end_row > start_row:
                ws.merge_cells(start_row=start_row, start_column=col, end_row=end_row, end_column=col)
                cell = ws.cell(row=start_row, column=col)
                cell.alignment = cell.alignment.copy(vertical='center')
            start_row = end_row + 1


def main():
    word_path = r'/Users/andresalfarofernandez/DocumentosPC/VisualStudio_code/Scripts/Autobastion/Templates/CIS_Microsoft_Windows_10_Enterprise_Benchmark_v3.0.0.docx'
    pdf_path = r'/Users/andresalfarofernandez/DocumentosPC/VisualStudio_code/Scripts/Autobastion/Templates/CIS_Microsoft_Windows_10_Enterprise_Benchmark_v3.0.0.pdf'
    excel_path = r'/Users/andresalfarofernandez/DocumentosPC/VisualStudio_code/Scripts/Autobastion/output.xlsx'

    with tqdm(total=100, desc="Procesando documento de Word a Excel", unit="porcentaje") as pbar:
        headings = extract_numbered_headings(word_path)
        pbar.update(20)
        titles1 = extract_controls(word_path, headings)
        pbar.update(10)

        titles2 = extract_domains(headings, titles1)

        pbar.set_description("Extrayendo textos de remediación")
        remediation_texts = extract_text_sections_pdf(pdf_path, "Remediation:")
        pbar.update(20)

        pbar.set_description("Extrayendo textos de valor por defecto")
        default_value_texts = extract_text_sections_word(word_path, "Default Value:")
        pbar.update(10)

        pbar.set_description("Extrayendo textos de verificación")
        verification_texts = extract_text_sections_pdf(pdf_path, "Audit:")
        pbar.update(20)

        pbar.set_description("Extrayendo textos de impacto")
        impact_texts = extract_text_sections_word(word_path, "Description:")
        pbar.update(15)

        pbar.set_description("Escribiendo datos en el archivo de Excel")
        write_titles_to_excel(titles2, remediation_texts, default_value_texts, verification_texts, impact_texts, excel_path)
        pbar.update(5)

    print(f"Datos copiados y pegados en {excel_path}")

if __name__ == "__main__":
    main()
