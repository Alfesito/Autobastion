from docx import Document
from openpyxl import Workbook
from tqdm import tqdm
import PyPDF2
import re
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

def extract_titles(doc_path, headings):
    doc = Document(doc_path)
    titles = []
    domain = None
    current_subdomain1 = None
    current_subdomain2 = None
    current_control = None
    remediation = None
    verification = None
    impact = None
    default_value = None
    is_sub2 = False

    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Document(element).paragraphs[0]
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split(' ')[1])
                text = para.text.strip()

                if level == 2:  # Dominio
                    is_sub2 = False
                    domain = text
                    current_subdomain1 = None
                    current_subdomain2 = None
                    current_control = None
                elif level == 3 and not is_sub2:  # Subdominio 1
                    current_subdomain1 = text
                    current_subdomain2 = None
                    current_control = None
                    is_sub2 = True
                elif level == 3 and is_sub2:  # Subdominio 2
                    current_subdomain2 = text
                    current_control = None
                elif level == 4:  # Control
                    current_control = text.split(' (')[0]  # Elimina todo lo que hay a la derecha de un ' ('
                    # Extraer el número de control
                    for key, value in headings.items():
                        value_words = value.replace(' ', '')
                        control_words = current_control.replace(' ', '')
                        if value_words == control_words:
                            control_number = key
                            break  # Salir del bucle una vez que se encuentra una coincidencia
                        # Ajustar la longitud de control_words si es mayor que value_words
                        elif len(control_words) > len(value_words):
                            control_words = control_words.ljust(len(value_words))
                            if control_words == value_words:
                                control_number = key
                                break
                        else:
                            control_number = None
                    titles.append((domain, current_subdomain1, current_subdomain2, control_number, current_control, remediation, default_value, verification, impact))
        elif isinstance(element, CT_Tbl):
            # Procesar tablas si es necesario
            pass

    return titles

def extract_text_sections(doc_path, section_title):
    doc = Document(doc_path)
    texts = []
    section_text = ""
    in_section = False
    in_default_value = False
    exclude_phrases = ["CIS Controls:", "MITRE ATT&CK Mappings:", "Audit:", "Remediation:", "References:", "Additional Information:", "Default Value:"]  # Lista de frases a excluir

    def process_paragraph(para):
        nonlocal section_text
        if para.style.name == 'List Paragraph':
            lines = para.text.split(' If ')
            for i, line in enumerate(lines):
                if i == 0:
                    section_text += '- ' + line.strip() + "\n"
                else:
                    section_text += "\nIf " + line.strip() + "\n"
        else:
            section_text += para.text.strip() + "\n"

    def process_table(table):
        nonlocal section_text
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Document(element).paragraphs[0]
            if section_title != "Default Value:":
                if para.style.name.startswith('Heading') and para.text in exclude_phrases and in_section:
                    texts.append(section_text.strip())
                    in_section = False
                    section_text = ""
                if para.style.name.startswith('Heading') and section_title == para.text:
                    in_section = True
                elif in_section and not any(exclude in para.text for exclude in exclude_phrases):
                    process_paragraph(para)
                elif para.text in exclude_phrases and in_section:
                    texts.append(section_text.strip())
                    in_section = False
                    section_text = ""
            else:
                if in_default_value and para.text in exclude_phrases and para.text != "Default Value:":
                    if not in_section:
                        section_text = "-"
                    texts.append(section_text.strip())
                    in_section = False
                    in_default_value = False
                    section_text = ""
                if para.style.name.startswith('Heading') and para.text == "Remediation:":
                    in_default_value = True
                elif para.style.name.startswith('Heading') and section_title == para.text and in_default_value:
                    in_section = True
                elif in_section:
                    process_paragraph(para)
        elif isinstance(element, CT_Tbl):
            if in_section:
                table = Document(element).tables[0]
                process_table(table)

    if in_section:
        texts.append(section_text.strip())

    return texts

def extract_numbered_headings(pdf_path):
    headings = {}
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text = page.extract_text()
            lines = text.split('\n')
            for line in lines:
                matches = re.findall(r'(\d+\.\d+(?:\.\d+)*)\s*\s*(.*)', line)
                if matches:
                    for match in matches:
                        clean_value = match[1].strip().replace(' -', '-').split(' (')[0].split(' .')[0]
                        if match[0] in headings:
                            break
                        else:
                            headings[match[0]] = clean_value
    return headings

def write_titles_to_excel(titles, remediation_texts, default_value_texts, verification_texts, impact_texts, excel_path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Controls"
    headers = ["Dominio", "Subdominio1", "Subdominio2", "ID", "Control", "Remediación", "Valor por defecto", "Verificación", "Impacto"]
    ws.append(headers)
    extended_titles = []

    for i in range(len(titles)):
        extended_title = list(titles[i])  # Convertir la tupla a lista
        if i < len(remediation_texts):
            extended_title[5] = remediation_texts[i]
        if i < len(default_value_texts):
            extended_title[6] = default_value_texts[i]
        if i < len(verification_texts):
            extended_title[7] = verification_texts[i]
        if i < len(impact_texts):
            extended_title[8] = impact_texts[i]
        extended_titles.append(extended_title)

    for i, extended_title in enumerate(tqdm(extended_titles, desc="Copiando datos a Excel", unit="fila")):
        ws.append(extended_title)
    
    merge_consecutive_rows(ws)
    wb.save(excel_path)

def merge_consecutive_rows(ws):
    max_row = ws.max_row
    for col in [1, 3]:
        start_row = 2
        while start_row <= max_row:
            end_row = start_row
            while end_row < max_row and ws.cell(row=end_row, column=col).value == ws.cell(row=end_row + 1, column=col).value and ws.cell(row=end_row + 1, column=col).value != None:
                end_row += 1
            if end_row > start_row:
                ws.merge_cells(start_row=start_row, start_column=col, end_row=end_row, end_column=col)
                cell = ws.cell(row=start_row, column=col)
                cell.alignment = cell.alignment.copy(vertical='center')
            start_row = end_row + 1

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []

    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Document(element).paragraphs[0]
            full_text.append(para.text)
        elif isinstance(element, CT_Tbl):
            table = Document(element).tables[0]
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text.append(para.text)

    return full_text

def main():
    word_path_en = r'D:\Usuarios\Ralfamole\Documentos\Cosas de Andres\VS\Autobastion\Templates\CIS_Debian_Linux_10_Benchmark_v2.0.0.docx'
    pdf_path = r'D:\Usuarios\Ralfamole\Documentos\Cosas de Andres\VS\Autobastion\Templates\CIS_Debian_Linux_10_Benchmark_v2.0.0.pdf'
    excel_path = r'D:\Usuarios\Ralfamole\Documentos\Cosas de Andres\VS\Autobastion\output.xlsx'

    with tqdm(total=100, desc="Procesando documento de Word a Excel", unit="porcentaje") as pbar:
        headings = extract_numbered_headings(pdf_path)
        pbar.update(20)
        titles = extract_titles(word_path_en, headings)
        pbar.update(10)

        pbar.set_description("Extrayendo textos de remediación")
        remediation_texts = extract_text_sections(word_path_en, "Remediation:")
        pbar.update(20)

        pbar.set_description("Extrayendo valores por defecto")
        default_value_texts = extract_text_sections(word_path_en, "Default Value:")
        pbar.update(20)

        pbar.set_description("Extrayendo textos de verificación")
        verification_texts = extract_text_sections(word_path_en, "Audit:")
        pbar.update(20)

        pbar.set_description("Extrayendo impactos")
        impact_texts = extract_text_sections(word_path_en, "Impact:")
        pbar.update(10)

        pbar.set_description("Escribiendo a Excel")
        write_titles_to_excel(titles, remediation_texts, default_value_texts, verification_texts, impact_texts, excel_path)
        pbar.update(20)

    print("Proceso completado.")

if __name__ == "__main__":
    main()
