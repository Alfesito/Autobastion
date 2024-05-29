from docx import Document
from openpyxl import Workbook
from tqdm import tqdm
import PyPDF2
import re

def extract_titles(doc_path, headings):
    doc = Document(doc_path)
    titles = []
    domain = None
    current_subdomain1 = None
    current_subdomain2 = None
    current_control = None
    remediation = None
    verification = None
    impact = None
    is_sub2 = False

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.split(' ')[1])
            text = para.text.strip()

            if level == 2:  # Dominio
                is_sub2 = False
                domain = text
                current_subdomain1 = None
                current_subdomain2 = None
                current_control = None
            elif level == 4 and not is_sub2:  # Subdominio 1
                current_subdomain1 = text
                current_subdomain2 = None
                current_control = None
                is_sub2 = True
            elif level == 4 and is_sub2:  # Subdominio 2
                current_subdomain2 = text
                current_control = None
            elif level == 3:  # Control
                current_control = text.split(' (')[0]  # Elimina todo lo que hay a la derecha de un ' ('
                # Extraer el número de control
                for key, value in headings.items():
                    value_words = value.replace(' ', '')
                    control_words = current_control.replace(' ', '')
                    if value_words == control_words:
                        control_number = key
                        break  # Salir del bucle una vez que se encuentra una coincidencia
                    # Ajustar la longitud de control_words si es mayor que value_words
                    elif len(control_words) > len(value_words):
                        control_words = control_words.ljust(len(value_words))
                        if control_words == value_words:
                            control_number = key
                            break
                    else:
                        control_number = None
                titles.append((domain, current_subdomain1, current_subdomain2, control_number, current_control, remediation, verification, impact))

    return titles

def extract_text_sections(doc_path, section_title):
    doc = Document(doc_path)
    texts = []
    section_text = ""
    in_section = False
    exclude_phrases = ["CIS Controls:", "MITRE ATT&CK Mappings:", "References:", "Default Value:"]  # Lista de frases a excluir
    include_phrases = ["-OR-", "OR", "-- OR --", "--OR--"] # Lista de frases a incluir

    def process_paragraph(para):
        nonlocal section_text
        if para.style.name == 'List Paragraph':
            lines = para.text.split(' If ')
            for i, line in enumerate(lines):
                if i == 0:
                    section_text += '- ' + line.strip() + "\n"
                else:
                    section_text += "\nIf " + line.strip() + "\n"
        else:
            section_text += para.text.strip() + "\n"

    for para in doc.paragraphs:
        if section_title == "Description:":
            if para.style.name.startswith('Heading') and in_section and (any(exclude in para.text for exclude in exclude_phrases) or not any(include in para.text for include in include_phrases)) and para.text == "Audit:":
                texts.append(section_text.strip())
                in_section = False
                section_text = ""
            if section_title in para.text:
                in_section = True
            elif in_section and not any(exclude in para.text for exclude in exclude_phrases):
                if para.text != "Rationale:":
                    process_paragraph(para)
            else:
                in_section = False
        else:
            if para.style.name.startswith('Heading') and in_section and (any(exclude in para.text for exclude in exclude_phrases) or not any(include in para.text for include in include_phrases)):
                texts.append(section_text.strip())
                in_section = False
                section_text = ""
            if section_title in para.text:
                in_section = True
            elif in_section and not any(exclude in para.text for exclude in exclude_phrases):
                process_paragraph(para)
            else:
                in_section = False

    if in_section:
        texts.append(section_text.strip())

    return texts

def extract_numbered_headings(pdf_path):
    headings = {}
    with open(pdf_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page in pdf_reader.pages:
            text = page.extract_text()
            lines = text.split('\n')
            for line in lines:
                matches = re.findall(r'(\d+\.\d+(?:\.\d+)*)\s*\s*(.*)', line)
                if matches:
                    for match in matches:
                        clean_value = match[1].strip().replace(' -', '-').split(' (')[0].split(' .')[0]
                        if match[0] in headings:
                            break
                        else:
                            headings[match[0]] = clean_value
    return headings

def write_titles_to_excel(titles, remediation_texts, verification_texts, impact_texts, extracted_text, excel_path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Controls"

    headers = ["Dominio", "Subdominio1", "Subdominio2", "ID", "Control", "Remediación", "Verificación", "Impacto"]
    ws.append(headers)

    extended_titles = []
    for i in range(len(titles)):
        extended_title = list(titles[i])
        if i < len(remediation_texts):
            extended_title[5] = remediation_texts[i]
        if i < len(verification_texts):
            extended_title[6] = verification_texts[i]
        if i < len(impact_texts):
            extended_title[7] = impact_texts[i]
        extended_titles.append(extended_title)

    for i, extended_title in enumerate(tqdm(extended_titles, desc="Copiando datos a Excel", unit="fila")):
        ws.append(extended_title)

    merge_consecutive_rows(ws)
    
    # Agregar el texto extraído del documento en una nueva hoja
    ws_text = wb.create_sheet("Extracted Text")
    for index, text in enumerate(extracted_text, start=1):
        ws_text.cell(row=index, column=1, value=text)
    
    wb.save(excel_path)

def merge_consecutive_rows(ws):
    max_row = ws.max_row
    for col in [1, 3]:
        start_row = 2
        while start_row <= max_row:
            end_row = start_row
            while end_row < max_row and ws.cell(row=end_row, column=col).value == ws.cell(row=end_row + 1, column=col).value and ws.cell(row=end_row + 1, column=col).value != None:
                end_row += 1
            if end_row > start_row:
                ws.merge_cells(start_row=start_row, start_column=col, end_row=end_row, end_column=col)
                cell = ws.cell(row=start_row, column=col)
                cell.alignment = cell.alignment.copy(vertical='center')
            start_row = end_row + 1

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for shape in doc.inline_shapes:
        if shape._inline.graphic.graphicData.textbox:
            textbox = shape._inline.graphic.graphicData.textbox
            for paragraph in textbox.content.children:
                if hasattr(paragraph, 'text'):
                    full_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text.append(para.text)

    return full_text

def main():
    word_path_en = r'C:\Users\aalfarofernandez\OneDrive - Deloitte (O365D)\Documents\Scripts\AutoBast\Templates\CIS_Debian_Linux_10_Benchmark_v2.0.0.docx'
    pdf_path = r'C:\Users\aalfarofernandez\OneDrive - Deloitte (O365D)\Documents\Scripts\AutoBast\Templates\CIS_Debian_Linux_10_Benchmark_v2.0.0.pdf'
    excel_path = r'C:\Users\aalfarofernandez\OneDrive - Deloitte (O365D)\Documents\Scripts\AutoBast\output.xlsx'

    with tqdm(total=100, desc="Procesando documento de Word a Excel", unit="porcentaje") as pbar:
        headings = extract_numbered_headings(pdf_path)
        pbar.update(20)
        titles = extract_titles(word_path_en, headings)
        pbar.update(10)

        pbar.set_description("Extrayendo textos de remediación")
        remediation_texts = extract_text_sections(word_path_en, "Remediation:")
        pbar.update(20)

        pbar.set_description("Extrayendo textos de verificación")
        verification_texts = extract_text_sections(word_path_en, "Audit:")
        pbar.update(20)

        pbar.set_description("Extrayendo textos de impacto")
        impact_texts = extract_text_sections(word_path_en, "Description:")  # Cambiado a "Description:" o "Descripción:"
        pbar.update(10)

        pbar.set_description("Extrayendo todo el texto del documento")
        extracted_text = extract_text_from_docx(word_path_en)
        pbar.update(10)

        pbar.set_description("Escribiendo datos en el archivo de Excel")
        write_titles_to_excel(titles, remediation_texts, verification_texts, impact_texts, extracted_text, excel_path)
        pbar.update(10)

    print(f"Datos copiados y pegados en {excel_path}")

if __name__ == "__main__":
    main()
